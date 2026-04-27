from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "posts.json"
OUTPUT_DIR = ROOT / "docs" / "guides"
QA_DIR = ROOT / "docs" / "qa"

ACCENT = "0F766E"
INK = "111827"
MUTED = "5C667A"
BLUE = "2556A3"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    heading.paragraph_format.space_after = Pt(4)
    for run in heading.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(INK if level == 1 else ACCENT)


def add_para(document: Document, text: str, color: str = INK, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.size = Pt(9.8)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_numbers(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.size = Pt(9.8)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(document: Document, title: str, body: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    title_run = paragraph.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(BLUE)
    body_run = paragraph.add_run(body)
    body_run.font.color.rgb = RGBColor.from_string(MUTED)


def add_lifecycle_block(document: Document, phase: str, purpose: str, actions: list[str]) -> None:
    add_heading(document, phase, 2)
    add_para(document, purpose, color=MUTED)
    add_bullets(document, actions)


def build_guide(guide: dict) -> Document:
    document = Document()
    configure_document(document)

    brand = document.add_paragraph()
    brand_run = brand.add_run("FusionPrep AI")
    brand_run.bold = True
